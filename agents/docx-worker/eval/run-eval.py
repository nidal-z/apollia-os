"""Runner d'eval pour docx-worker.

Crée des fixtures .docx temporaires (sample avec headings + table + list,
doc avec hyperlink, doc appendable, template Jinja2, .docm stub), exécute
les cas définis dans ``eval/cases.jsonl``, compare les résultats.

Usage :
    python3 eval/run-eval.py
    python3 eval/run-eval.py --case happy-write-with-styles
    python3 eval/run-eval.py --verbose
"""

from __future__ import annotations

import argparse
import asyncio
import importlib.util
import json
import shutil
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any


# ─── Case structure ────────────────────────────────────────────────────────


@dataclass
class EvalCase:
    name: str
    payload: dict[str, Any]
    expected: dict[str, Any]
    skill_id: str | None = None
    note: str = ""


def load_cases(path: Path) -> list[EvalCase]:
    cases: list[EvalCase] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        raw = json.loads(line)
        cases.append(
            EvalCase(
                name=raw["name"],
                payload=raw["payload"],
                expected=raw["expected"],

                skill_id=raw.get("skill_id"),
                note=raw.get("note", ""),
            )
        )
    return cases


def render_placeholders(value: Any, fx_base: str) -> Any:
    if isinstance(value, str):
        return value.replace("${FX}", fx_base)
    if isinstance(value, list):
        return [render_placeholders(v, fx_base) for v in value]
    if isinstance(value, dict):
        return {k: render_placeholders(v, fx_base) for k, v in value.items()}
    return value


# ─── Fixture builder ───────────────────────────────────────────────────────


def build_fixtures(base: Path) -> None:
    """Create test fixtures :

    - ``sample.docx``         : heading + paragraphs + table 2x3 + bulleted list
    - ``with-hyperlink.docx`` : doc with a hyperlink (built via XML insertion)
    - ``appendable.docx``     : copy of sample for append test
    - ``appendable2.docx``    : another copy for append+new_section test
    - ``template.docx``       : Jinja2 template with placeholders
    - ``macros.docm``         : empty stub for UNSUPPORTED_FORMAT test
    """
    base.mkdir(parents=True, exist_ok=True)

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # sample.docx
    sample = Document()
    sample.add_heading("Titre 1", level=1)
    sample.add_paragraph("Paragraphe normal.")
    sample.add_heading("Titre 2", level=2)
    sample.add_paragraph("Autre paragraphe.")
    tbl = sample.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "h1"
    tbl.cell(0, 1).text = "h2"
    tbl.cell(0, 2).text = "h3"
    tbl.cell(1, 0).text = "a"
    tbl.cell(1, 1).text = "b"
    tbl.cell(1, 2).text = "c"
    sample.add_paragraph("Item bullet 1", style="List Bullet")
    sample.add_paragraph("Item bullet 2", style="List Bullet")
    sample_path = base / "sample.docx"
    sample.save(sample_path)

    # with-hyperlink.docx — insert a hyperlink via XML
    hl_doc = Document()
    p = hl_doc.add_paragraph("Visit ")
    # Build hyperlink XML : <w:hyperlink r:id="rId1"><w:r><w:t>Apollia</w:t></w:r></w:hyperlink>
    part = p.part
    r_id = part.relate_to(
        "https://apollia.fr",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Apollia"
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)
    p.add_run(" website.")
    hl_doc.save(base / "with-hyperlink.docx")

    # Copies for append
    shutil.copy(sample_path, base / "appendable.docx")
    shutil.copy(sample_path, base / "appendable2.docx")

    # template.docx — Jinja2 placeholders
    tpl = Document()
    tpl.add_heading("Devis", level=1)
    tpl.add_paragraph("Client : {{ client_name }}")
    tpl.add_paragraph("Montant : {{ amount }} €")
    tpl.add_paragraph("Détail :")
    tpl.add_paragraph("{% for item in items %}- {{ item.label }} : {{ item.price }} €{% endfor %}")
    tpl.save(base / "template.docx")

    # macros.docm — stub (rejected on extension)
    (base / "macros.docm").write_bytes(b"PK\x03\x04dummy-macros")


# ─── MockCtx ───────────────────────────────────────────────────────────────


class MockCtx:
    """Minimal RuntimeContext stub — sufficient for the deterministic worker."""

    llm = None
    workspace = None

    class _Tools:
        async def call(self, name: str, args: dict[str, Any]) -> dict[str, Any]:
            return {"error": {"code": "no_tool_in_eval", "message": f"{name} not mocked"}}

    tools = _Tools()

    def log(self, level: str, msg: str) -> None:
        if level in ("warn", "error"):
            print(f"  [worker.{level}] {msg}", file=sys.stderr)


# ─── Case runner ───────────────────────────────────────────────────────────


def _build_task(payload: dict[str, Any], skill_id: str | None) -> dict[str, Any]:
    """Build a minimal AIPTask. ``skill_id`` is propagated as the runtime would."""
    task: dict[str, Any] = {
        "task_id": "eval-task",
        "context_id": "eval-context",
        "input": {"parts": [{"type": "data", "data": payload}]},
    }
    if skill_id is not None:
        task["skill_id"] = skill_id
    return task


def _parse_output(result: dict[str, Any]) -> dict[str, Any] | None:
    try:
        return json.loads(result["output"][0]["text"])
    except (KeyError, IndexError, json.JSONDecodeError):
        return None


async def run_case(agent: Any, case: EvalCase, fx_base: str, verbose: bool) -> bool:
    payload = render_placeholders(case.payload, fx_base)
    task = _build_task(payload, case.skill_id)
    result = await agent.run(task, MockCtx())

    actual_status = result.get("status")
    expected_status = case.expected.get("status")
    if actual_status != expected_status:
        print(f"  ❌ status mismatch — expected {expected_status}, got {actual_status}")
        if verbose:
            print(f"     full result: {result}")
        return False

    if expected_status == "failed":
        expected_code = case.expected.get("code")
        actual_code = result.get("error", {}).get("code")
        if expected_code and actual_code != expected_code:
            print(f"  ❌ error code mismatch — expected {expected_code}, got {actual_code}")
            if verbose:
                print(f"     full result: {result}")
            return False

    if expected_status == "completed":
        output = _parse_output(result)
        if output is None:
            print(f"  ❌ output not parseable JSON: {result}")
            return False

        keys_expected = case.expected.get("schema_match")
        if keys_expected:
            missing = [k for k in keys_expected if k not in output]
            if missing:
                print(f"  ❌ missing keys in output : {missing}")
                if verbose:
                    print(f"     output: {output}")
                return False

    print("  ✅")
    return True


# ─── Main ──────────────────────────────────────────────────────────────────


def _load_worker_module() -> Any:
    parent = Path(__file__).resolve().parent.parent
    worker_path = parent / "docx-worker.py"
    spec = importlib.util.spec_from_file_location("docx_worker_under_test", str(worker_path))
    if spec is None or spec.loader is None:
        raise RuntimeError(f"unable to import worker from {worker_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules["docx_worker_under_test"] = module
    spec.loader.exec_module(module)
    return module


async def main() -> int:
    parser = argparse.ArgumentParser(description="Eval runner for docx-worker")
    parser.add_argument("--case", help="Run only this named case")
    parser.add_argument("--verbose", action="store_true", help="Print full results on fail")
    parser.add_argument(
        "--cases-file",
        default=str(Path(__file__).parent / "cases.jsonl"),
        help="Path to JSONL cases file",
    )
    parser.add_argument("--keep-fixtures", action="store_true", help="Keep fixture dir after run")
    args = parser.parse_args()

    worker_module = _load_worker_module()
    agent = worker_module.agent

    fx_dir = Path(tempfile.mkdtemp(prefix="docx-worker-eval-"))
    print(f"Building fixtures in {fx_dir}")
    build_fixtures(fx_dir)
    fx_base = str(fx_dir)

    try:
        cases = load_cases(Path(args.cases_file))
        if args.case:
            cases = [c for c in cases if c.name == args.case]
            if not cases:
                print(f"No case named {args.case!r}")
                return 1

        print(f"\nRunning {len(cases)} case(s) against docx-worker…")
        pass_count = 0
        for case in cases:
            print(f"\n• {case.name}")
            if case.note:
                print(f"  {case.note}")
            ok = await run_case(agent, case, fx_base, args.verbose)
            if ok:
                pass_count += 1

        total = len(cases)
        ratio = 100 * pass_count // total if total else 0
        print(f"\n{'=' * 60}\n{pass_count}/{total} cases passed ({ratio}%)")
        return 0 if pass_count == total else 1
    finally:
        if not args.keep_fixtures:
            shutil.rmtree(fx_dir, ignore_errors=True)
        else:
            print(f"\nFixtures kept at : {fx_dir}")


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
