"""docx-worker — Manipulation de fichiers Word (.docx).

Worker Apollia OS standalone qui isole la manipulation .docx du LLM director.
Cinq skills A2A déterministes :

- ``docx.read``                 — lecture structurée (blocs + sections + hyperlinks).
- ``docx.write``                — création avec styles avancés (named_styles, headers/footers).
- ``docx.append-section``       — ajout de blocs en fin de document.
- ``docx.extract-tables``       — extraction ciblée des tables uniquement.
- ``docx.render-from-template`` — rendu Jinja2 via docxtpl (mode strict par défaut).

Dispatch multi-skills : le SDK route ``task.skill_id`` vers la méthode
``@skill`` correspondante (ex. ``docx.read``). Les paramètres sont
extraits depuis la signature typée.

Headers/footers — 10 placeholders disponibles :
- ``{page}``, ``{total_pages}``  — champs Word PAGE / NUMPAGES (dynamiques)
- ``{date}``, ``{time}``, ``{datetime}``  — date/time du moment d'écriture
  (formatables : ``{date:%d/%m/%Y}``)
- ``{author}``, ``{title}``, ``{subject}``  — depuis core_properties
- ``{filename}``  — basename du output_path
- ``{section_num}``  — index 1-based de la section
"""

from __future__ import annotations

import datetime
import json
import re
from pathlib import Path
from typing import Annotated, Any

from apollia import DomainError, agent, skill
from apollia.types import Ctx

# Resolve ``schemas`` from the worker dir even when another agent has
# previously populated ``sys.modules['schemas']`` from its own folder.
import sys as _sys
from pathlib import Path as _PathForBootstrap
_WORKER_DIR = str(_PathForBootstrap(__file__).resolve().parent)
if _WORKER_DIR in _sys.path:
    _sys.path.remove(_WORKER_DIR)
_sys.path.insert(0, _WORKER_DIR)
_sys.modules.pop("schemas", None)

# Canonical input TypedDicts — structural input_schema for mid-market LLMs.
# DocxBlock is a polymorphic union (discriminated by ``type``) — rendered as
# ``anyOf`` in the JSON Schema so the LLM sees every block variant.
from schemas import (  # type: ignore[import-not-found]  # noqa: E402
    DocxBlock,
    DocumentSetup,
    SectionWriteSpec,
    StyleSpec,
)

MAX_FILE_BYTES: int = 100 * 1024 * 1024  # 100 MiB
PAPER_SIZES: tuple[str, ...] = ("A4", "Letter", "Legal", "A3", "A5")
SUPPORTED_IMAGE_EXTS: tuple[str, ...] = (".png", ".jpg", ".jpeg", ".gif", ".tiff", ".tif", ".bmp")

# Paper size dimensions (cm) — (width, height) portrait orientation
PAPER_DIMS_CM: dict[str, tuple[float, float]] = {
    "A4":     (21.0, 29.7),
    "Letter": (21.59, 27.94),
    "Legal":  (21.59, 35.56),
    "A3":     (29.7, 42.0),
    "A5":     (14.8, 21.0),
}

# Placeholder field codes for Word fields (dynamic)
FIELD_PLACEHOLDERS: dict[str, str] = {"page": "PAGE", "total_pages": "NUMPAGES"}


class _DocxError(DomainError):
    """Legacy alias kept for the internal ``_do_*`` helpers.

    They raise ``_DocxError`` which is itself a :class:`DomainError`, so the
    dispatch boundary handles it as a typed failure with stable code +
    message + details. Future cleanup may inline-replace these.
    """


@agent(
    name="docx-worker",
    version="0.1.0",
    description=(
        "Manipulation de fichiers Word (.docx) : lire, écrire, ajouter, "
        "extraire des tables, rendre des templates Jinja2 — préserve "
        "styles, sections, headers/footers."
    ),
    tags=("docx", "word", "office", "python-docx", "docxtpl", "template", "worker"),
    agent_type="worker",
    step_budget={"max_steps": 1, "max_tool_calls": 5, "wall_clock_secs": 300},
)
class DocxWorker:
    """5 skills .docx déterministes (decorator-first)."""

    @skill(
        "docx.read",
        description=(
            "Read a .docx and return structured JSON: ordered blocks, "
            "section metadata, hyperlinks, and used styles."
        ),
        examples=[
            {"path": "/path/to/document.docx", "include_runs": False},
        ],
    )
    async def read_docx(
        self,
        path: Annotated[str, "Filesystem path to the source .docx file."],
        include_runs: Annotated[
            bool,
            "Include per-run formatting (bold/italic/font/color) inside paragraphs. Larger payload.",
        ] = False,
        include_styles: Annotated[
            bool,
            "Include the list of paragraph/run styles used in the document.",
        ] = True,
        max_blocks: Annotated[
            int,
            "Hard cap on top-level blocks returned (truncates with truncated=true).",
        ] = 50_000,
        ctx: Ctx = None,
    ) -> dict[str, Any]:
        """Lecture structurée."""
        return _do_read(
            {
                "path": path,
                "include_runs": include_runs,
                "include_styles": include_styles,
                "max_blocks": max_blocks,
            }
        )

    @skill(
        "docx.write",
        description=(
            "Create a multi-section .docx from a structured spec (sections, "
            "ordered blocks, named styles, headers/footers, page setup)."
        ),
        examples=[
            {
                "output_path": "/tmp/report.docx",
                "sections": [
                    {
                        "paper_size": "A4",
                        "orientation": "portrait",
                        "blocks": [
                            {"type": "heading", "text": "Quarterly Report", "level": 1},
                            {"type": "paragraph", "text": "Executive summary..."},
                            {
                                "type": "table",
                                "rows": [
                                    ["Product", "Q1", "Q2"],
                                    ["Widget", "100", "120"],
                                ],
                            },
                            {"type": "page_break"},
                        ],
                    },
                ],
                "named_styles": {
                    "emphasis": {"font": {"bold": True, "color": "#1F4E78"}},
                },
            },
        ],
    )
    async def write_docx(
        self,
        output_path: Annotated[str, "Filesystem path of the .docx to create."],
        sections: list[SectionWriteSpec],
        named_styles: dict[str, StyleSpec] | None = None,
        document_setup: DocumentSetup | None = None,
        overwrite: Annotated[bool, "Allow overwriting an existing output file."] = False,
        ctx: Ctx = None,
    ) -> dict[str, Any]:
        """Crée un .docx."""
        return _do_write(
            {
                "output_path": output_path,
                "sections": sections,
                "named_styles": named_styles,
                "document_setup": document_setup,
                "overwrite": overwrite,
            }
        )

    @skill(
        "docx.append_section",
        description=(
            "Append blocks to an existing .docx without touching its prior "
            "content, styles, headers, footers, or sections."
        ),
        examples=[
            {
                "path": "/tmp/report.docx",
                "blocks": [
                    {"type": "heading", "text": "Appendix", "level": 1},
                    {"type": "paragraph", "text": "Additional details follow..."},
                ],
                "section_break_before": True,
            },
        ],
    )
    async def append_section(
        self,
        path: Annotated[str, "Filesystem path to the existing .docx file."],
        blocks: list[DocxBlock],
        named_styles: dict[str, StyleSpec] | None = None,
        section_break_before: Annotated[
            bool,
            "Insert a new section break before appending the blocks.",
        ] = False,
        new_section_setup: SectionWriteSpec | None = None,
        ctx: Ctx = None,
    ) -> dict[str, Any]:
        """Ajoute des blocs à un .docx."""
        return _do_append_section(
            {
                "path": path,
                "blocks": blocks,
                "named_styles": named_styles,
                "section_break_before": section_break_before,
                "new_section_setup": new_section_setup,
            }
        )

    @skill(
        "docx.extract_tables",
        description=(
            "Extract only the tables from a .docx. Faster than docx.read "
            "when the caller doesn't need paragraphs or styles."
        ),
        examples=[
            {"path": "/path/to/document.docx"},
        ],
    )
    async def extract_tables(
        self,
        path: Annotated[str, "Filesystem path to the source .docx file."],
        include_headers: Annotated[
            bool,
            "Expose the first row of each table as 'headers' for convenience.",
        ] = True,
        ctx: Ctx = None,
    ) -> dict[str, Any]:
        """Extraction tables."""
        return _do_extract_tables({"path": path, "include_headers": include_headers})

    @skill(
        "docx.render_from_template",
        description=(
            "Render a .docx template containing Jinja2 placeholders via "
            "docxtpl, injecting the provided context. Strict mode by default."
        ),
        examples=[
            {
                "template_path": "/path/to/invoice_template.docx",
                "output_path": "/tmp/invoice_42.docx",
                "context": {
                    "client_name": "Acme Corp",
                    "invoice_number": "INV-2026-042",
                    "total": 1250.00,
                    "items": [
                        {"description": "Consulting", "amount": 1000.00},
                        {"description": "Support", "amount": 250.00},
                    ],
                },
            },
        ],
    )
    async def render_from_template(
        self,
        template_path: Annotated[str, "Filesystem path to the source .docx template with Jinja2 placeholders."],
        output_path: Annotated[str, "Filesystem path of the rendered .docx to create."],
        context: Annotated[
            dict[str, Any],
            "Jinja2 context: arbitrary keys consumed by the template. Values can be strings, numbers, lists, nested dicts.",
        ],
        strict_undefined: Annotated[
            bool,
            "When true, missing placeholders raise an error instead of rendering as empty.",
        ] = True,
        overwrite: Annotated[bool, "Allow overwriting an existing output file."] = False,
        ctx: Ctx = None,
    ) -> dict[str, Any]:
        """Rendu Jinja2 .docx."""
        return _do_render_from_template(
            {
                "template_path": template_path,
                "output_path": output_path,
                "context": context,
                "strict_undefined": strict_undefined,
                "overwrite": overwrite,
            }
        )


# ─── Operation: read ───────────────────────────────────────────────────────


def _do_read(payload: dict[str, Any]) -> dict[str, Any]:
    path = _require_str(payload, "path")
    _validate_docx_path(path, must_exist=True)

    include_runs = bool(payload.get("include_runs", False))
    include_styles = bool(payload.get("include_styles", True))
    max_blocks = int(payload.get("max_blocks", 50_000))

    from docx import Document

    try:
        doc = Document(path)
    except Exception as exc:
        raise _DocxError("PARSE_ERROR", f"Impossible de lire le .docx : {exc}") from exc

    # Sections metadata (ordered)
    sections_out: list[dict[str, Any]] = []
    for section in doc.sections:
        sections_out.append({
            "orientation": _section_orientation(section),
            "margins_cm": {
                "top": _emu_to_cm(section.top_margin),
                "bottom": _emu_to_cm(section.bottom_margin),
                "left": _emu_to_cm(section.left_margin),
                "right": _emu_to_cm(section.right_margin),
            },
            "paper_size": _detect_paper_size(section),
            "header_text": _extract_header_footer_text(section.header),
            "footer_text": _extract_header_footer_text(section.footer),
            "blocks": [],
        })

    # Iterate body children in order. All blocks → first section in v0.1.0
    # (mapping blocks to sections requires parsing sectPr elements — future).
    styles_used: set[str] = set()
    hyperlinks: list[dict[str, str]] = []
    block_count = 0
    truncated = False

    target_blocks = sections_out[0]["blocks"] if sections_out else []

    body = doc.element.body
    for elem in body.iterchildren():
        if block_count >= max_blocks:
            truncated = True
            break
        tag = elem.tag.rsplit("}", 1)[-1]
        if tag == "p":
            block = _read_paragraph(elem, doc, include_runs, include_styles, hyperlinks, styles_used)
            if block is not None:
                target_blocks.append(block)
                block_count += 1
        elif tag == "tbl":
            block = _read_table(elem, doc)
            target_blocks.append(block)
            block_count += 1
        # ignore sectPr and other top-level elements

    return {
        "path": path,
        "sections": sections_out,
        "styles_used": sorted(styles_used),
        "hyperlinks": hyperlinks,
        "total_blocks": block_count,
        "truncated": truncated,
    }


def _section_orientation(section: Any) -> str:
    from docx.enum.section import WD_ORIENT

    return "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"


def _detect_paper_size(section: Any) -> str | None:
    w = _emu_to_cm(section.page_width)
    h = _emu_to_cm(section.page_height)
    # Normalize portrait orientation for comparison
    width, height = (min(w, h), max(w, h))
    for name, (pw, ph) in PAPER_DIMS_CM.items():
        if abs(width - pw) < 0.2 and abs(height - ph) < 0.2:
            return name
    return None


def _extract_header_footer_text(hf: Any) -> str:
    return "\n".join(p.text for p in hf.paragraphs if p.text).strip() or None


def _read_paragraph(
    p_element: Any,
    doc: Any,
    include_runs: bool,
    include_styles: bool,
    hyperlinks: list[dict[str, str]],
    styles_used: set[str],
) -> dict[str, Any] | None:
    from docx.text.paragraph import Paragraph

    p = Paragraph(p_element, doc.part)
    style_name = (p.style.name if p.style else "Normal") or "Normal"
    if include_styles:
        styles_used.add(style_name)

    # Detect page break inside the paragraph (w:br w:type="page")
    has_page_break = False
    for run in p.runs:
        for br in run._r.findall(_qn("w:br")):
            if br.get(_qn("w:type")) == "page":
                has_page_break = True
                break
        if has_page_break:
            break

    # Detect heading
    if style_name.startswith("Heading "):
        try:
            level = int(style_name.split()[-1])
        except (ValueError, IndexError):
            level = 1
        block: dict[str, Any] = {"type": "heading", "text": p.text, "level": level}
        if include_styles:
            block["style_name"] = style_name
        return block

    if style_name == "Title":
        block = {"type": "heading", "text": p.text, "level": 1}
        if include_styles:
            block["style_name"] = style_name
        return block

    # Detect list (style "List Bullet" / "List Number" or numPr present)
    is_bullet = "List Bullet" in style_name
    is_number = "List Number" in style_name
    if is_bullet or is_number:
        block = {
            "type": "list",
            "text": p.text,
            "list_type": "bullet" if is_bullet else "number",
        }
        if include_styles:
            block["style_name"] = style_name
        return block

    # Detect hyperlinks
    for hl in p_element.findall(_qn("w:hyperlink")):
        rid = hl.get(_qn("r:id"))
        text_parts = [t.text for t in hl.iter(_qn("w:t")) if t.text]
        text = "".join(text_parts)
        if rid:
            try:
                rel = doc.part.rels[rid]
                hyperlinks.append({"text": text, "url": rel.target_ref})
            except KeyError:
                pass

    # Plain paragraph
    block = {"type": "paragraph", "text": p.text}
    if include_styles:
        block["style_name"] = style_name

    # Alignment
    align = p.alignment
    if align is not None:
        block["alignment"] = _alignment_to_str(align)

    if include_runs:
        block["runs"] = [_run_to_dict(r) for r in p.runs if r.text]

    if has_page_break and not p.text:
        # Pure page break paragraph
        return {"type": "page_break"}

    return block


def _run_to_dict(run: Any) -> dict[str, Any]:
    font = run.font
    out: dict[str, Any] = {"text": run.text}
    f: dict[str, Any] = {}
    if font.name:
        f["name"] = font.name
    if font.size is not None:
        f["size_pt"] = float(font.size.pt)
    if font.bold is not None:
        f["bold"] = bool(font.bold)
    if font.italic is not None:
        f["italic"] = bool(font.italic)
    if font.underline is not None:
        f["underline"] = bool(font.underline)
    if font.strike is not None:
        f["strike"] = bool(font.strike)
    if font.color and font.color.rgb is not None:
        f["color"] = str(font.color.rgb)
    if f:
        out["font"] = f
    return out


def _read_table(tbl_element: Any, doc: Any) -> dict[str, Any]:
    from docx.table import Table

    tbl = Table(tbl_element, doc.part)
    rows: list[list[str]] = []
    for row in tbl.rows:
        rows.append([cell.text for cell in row.cells])
    return {
        "type": "table",
        "rows": rows,
        "rows_count": len(rows),
        "cols_count": len(rows[0]) if rows else 0,
    }


# ─── Operation: write ──────────────────────────────────────────────────────


def _do_write(payload: dict[str, Any]) -> dict[str, Any]:
    output_path = _require_str(payload, "output_path")
    _validate_docx_extension(output_path)
    overwrite = bool(payload.get("overwrite", False))

    out = Path(output_path).expanduser().resolve()
    if out.exists() and not overwrite:
        raise _DocxError(
            "OUTPUT_EXISTS",
            f"Fichier de sortie existe déjà (overwrite=false) : {out}",
            details={"output_path": str(out)},
        )
    out.parent.mkdir(parents=True, exist_ok=True)

    sections_spec = payload.get("sections")
    if not isinstance(sections_spec, list) or not sections_spec:
        raise _DocxError(
            "MISSING_FIELD",
            "Champ 'sections' requis (liste non vide)",
            details={"field": "sections"},
        )

    named_styles = payload.get("named_styles") or {}
    document_setup = payload.get("document_setup") or {}

    if not isinstance(named_styles, dict):
        raise _DocxError("INVALID_TYPE", "named_styles doit être un objet", details={"field": "named_styles"})

    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()

    # Remove default empty paragraph if it exists
    body = doc.element.body
    for p in list(body.findall(_qn("w:p"))):
        body.remove(p)

    blocks_count = 0
    styles_applied_count = 0

    for sec_idx, section_spec in enumerate(sections_spec):
        if not isinstance(section_spec, dict):
            raise _DocxError(
                "INVALID_TYPE",
                f"sections[{sec_idx}] doit être un objet",
                details={"field": f"sections[{sec_idx}]"},
            )

        # Get or add section
        if sec_idx == 0:
            section = doc.sections[0]
        else:
            # Add new section via section break in last paragraph
            new_p = doc.add_paragraph()
            run = new_p.add_run()
            run.add_break(WD_BREAK.PAGE)
            section = doc.add_section()

        # Apply section setup (merged with document_setup defaults)
        _apply_section_setup(section, document_setup, section_spec)

        # Setup headers/footers
        ctx_dict = _build_ctx_dict(doc, output_path, sec_idx + 1)
        if section_spec.get("header"):
            _setup_header_footer(section, section_spec["header"], ctx_dict, is_header=True)
        if section_spec.get("footer"):
            _setup_header_footer(section, section_spec["footer"], ctx_dict, is_header=False)

        # Render blocks
        blocks = section_spec.get("blocks") or []
        if not isinstance(blocks, list):
            raise _DocxError(
                "INVALID_TYPE",
                f"sections[{sec_idx}].blocks doit être une liste",
                details={"field": f"sections[{sec_idx}].blocks"},
            )
        for block_idx, block in enumerate(blocks):
            if not isinstance(block, dict):
                raise _DocxError(
                    "INVALID_BLOCK",
                    f"sections[{sec_idx}].blocks[{block_idx}] doit être un objet",
                    details={"field": f"sections[{sec_idx}].blocks[{block_idx}]"},
                )
            styles_applied_count += _add_block(doc, block, named_styles)
            blocks_count += 1

    try:
        doc.save(str(out))
    except Exception as exc:
        raise _DocxError("EXECUTION_FAILED", f"Sauvegarde impossible : {exc}") from exc

    return {
        "output_path": str(out),
        "sections_count": len(sections_spec),
        "blocks_count": blocks_count,
        "styles_applied_count": styles_applied_count,
    }


def _apply_section_setup(section: Any, document_setup: dict[str, Any], section_spec: dict[str, Any]) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    # Merged setup: section overrides document defaults
    merged: dict[str, Any] = dict(document_setup)
    for key in ("orientation", "margins_cm", "paper_size"):
        if key in section_spec:
            merged[key] = section_spec[key]

    paper_size = merged.get("paper_size")
    if paper_size:
        if paper_size not in PAPER_SIZES:
            raise _DocxError(
                "INVALID_TYPE",
                f"paper_size invalide : {paper_size!r}. Valeurs : {list(PAPER_SIZES)}",
                details={"field": "paper_size", "got": paper_size},
            )
        w_cm, h_cm = PAPER_DIMS_CM[paper_size]
        section.page_width = Cm(w_cm)
        section.page_height = Cm(h_cm)

    orientation = merged.get("orientation")
    if orientation:
        if orientation not in ("portrait", "landscape"):
            raise _DocxError(
                "INVALID_TYPE",
                f"orientation invalide : {orientation!r}",
                details={"field": "orientation"},
            )
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap dimensions if currently portrait
            w, h = section.page_width, section.page_height
            if h > w:
                section.page_width, section.page_height = h, w
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            w, h = section.page_width, section.page_height
            if w > h:
                section.page_width, section.page_height = h, w

    margins = merged.get("margins_cm")
    if margins:
        if not isinstance(margins, dict):
            raise _DocxError("INVALID_TYPE", "margins_cm doit être un objet", details={"field": "margins_cm"})
        if "top" in margins:
            section.top_margin = Cm(float(margins["top"]))
        if "bottom" in margins:
            section.bottom_margin = Cm(float(margins["bottom"]))
        if "left" in margins:
            section.left_margin = Cm(float(margins["left"]))
        if "right" in margins:
            section.right_margin = Cm(float(margins["right"]))


def _build_ctx_dict(doc: Any, output_path: str, section_num: int) -> dict[str, str]:
    cp = doc.core_properties
    return {
        "author": cp.author or "",
        "title": cp.title or "",
        "subject": cp.subject or "",
        "filename": Path(output_path).name,
        "section_num": str(section_num),
    }


def _setup_header_footer(
    section: Any,
    hf_spec: Any,
    ctx_dict: dict[str, str],
    is_header: bool,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

    if not isinstance(hf_spec, dict):
        raise _DocxError(
            "INVALID_TYPE",
            "header/footer doit être un objet {left, center, right}",
            details={"field": "header_or_footer"},
        )

    hf = section.header if is_header else section.footer
    # Use the first paragraph (or create one)
    if hf.paragraphs:
        p = hf.paragraphs[0]
    else:
        p = hf.add_paragraph()

    # Clear existing content
    for r in list(p._p.findall(_qn("w:r"))):
        p._p.remove(r)

    left = str(hf_spec.get("left", ""))
    center = str(hf_spec.get("center", ""))
    right = str(hf_spec.get("right", ""))

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if left:
        _add_text_with_placeholders(p, left, ctx_dict)

    if center:
        p.add_run().add_tab()
        _add_text_with_placeholders(p, center, ctx_dict)

    if right:
        p.add_run().add_tab()
        _add_text_with_placeholders(p, right, ctx_dict)

    # Define tab stops: center and right
    tab_stops = p.paragraph_format.tab_stops
    # Clear existing tab stops
    try:
        tab_stops.clear_all()
    except Exception:
        pass

    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    if page_width is None or left_margin is None or right_margin is None:
        return
    from docx.shared import Emu

    usable = page_width - left_margin - right_margin
    tab_stops.add_tab_stop(Emu(usable // 2), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Emu(usable), WD_TAB_ALIGNMENT.RIGHT)


_PLACEHOLDER_RE = re.compile(r"\{([a-zA-Z_]+)(?::([^{}]+))?\}")


def _add_text_with_placeholders(paragraph: Any, text: str, ctx_dict: dict[str, str]) -> None:
    last_end = 0
    for match in _PLACEHOLDER_RE.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        key = match.group(1)
        fmt = match.group(2)

        if key in FIELD_PLACEHOLDERS:
            _insert_word_field(paragraph, FIELD_PLACEHOLDERS[key])
        elif key == "date":
            today = datetime.date.today()
            paragraph.add_run(today.strftime(fmt) if fmt else today.isoformat())
        elif key == "time":
            now = datetime.datetime.now()
            paragraph.add_run(now.strftime(fmt) if fmt else now.strftime("%H:%M"))
        elif key == "datetime":
            now = datetime.datetime.now()
            paragraph.add_run(now.strftime(fmt) if fmt else now.isoformat(timespec="seconds"))
        elif key in ctx_dict:
            paragraph.add_run(ctx_dict[key])
        else:
            # Unknown placeholder — keep literal
            paragraph.add_run(match.group(0))

        last_end = match.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _insert_word_field(paragraph: Any, field_code: str) -> None:
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(_qn("w:fldCharType"), "begin")
    r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(_qn("w:fldCharType"), "end")
    r.append(fld_end)


# ─── Block rendering ───────────────────────────────────────────────────────


def _add_block(doc: Any, block: dict[str, Any], named_styles: dict[str, Any]) -> int:
    """Add a block to the document. Returns the number of styles applied."""
    btype = block.get("type")
    if not btype:
        raise _DocxError(
            "INVALID_BLOCK",
            "Block sans 'type'",
            details={"block": block},
        )

    if btype == "paragraph":
        return _add_paragraph_block(doc, block, named_styles)
    if btype == "heading":
        return _add_heading_block(doc, block, named_styles)
    if btype == "table":
        return _add_table_block(doc, block, named_styles)
    if btype == "list":
        return _add_list_block(doc, block, named_styles)
    if btype == "image":
        return _add_image_block(doc, block)
    if btype == "page_break":
        _add_page_break(doc)
        return 0
    if btype == "section_break":
        # Section breaks add new sections — handled at section level
        # In a block context, treat as page_break for v0.1.0
        _add_page_break(doc)
        return 0

    raise _DocxError(
        "INVALID_BLOCK",
        f"Type de block non reconnu : {btype!r}",
        details={"got": btype},
    )


def _add_paragraph_block(doc: Any, block: dict[str, Any], named_styles: dict[str, Any]) -> int:
    p = doc.add_paragraph()
    style_id = block.get("style")
    styles_applied = 0

    text = block.get("text")
    runs_spec = block.get("runs")

    if runs_spec:
        if not isinstance(runs_spec, list):
            raise _DocxError("INVALID_BLOCK", "paragraph.runs doit être une liste", details={"field": "runs"})
        for run_spec in runs_spec:
            if not isinstance(run_spec, dict):
                raise _DocxError("INVALID_BLOCK", "paragraph.runs[i] doit être un objet", details={"field": "runs[i]"})
            run = p.add_run(str(run_spec.get("text", "")))
            font_spec = run_spec.get("font")
            if font_spec:
                _apply_font_to_run(run, font_spec)
    elif text is not None:
        p.add_run(str(text))

    alignment = block.get("alignment")
    if alignment:
        p.alignment = _str_to_alignment(alignment)

    if style_id:
        if style_id not in named_styles:
            raise _DocxError(
                "STYLE_NOT_FOUND",
                f"Style '{style_id}' référencé mais absent de named_styles",
                details={"style_id": style_id, "available": list(named_styles)},
            )
        _apply_paragraph_style(p, named_styles[style_id])
        styles_applied += 1

    return styles_applied


def _add_heading_block(doc: Any, block: dict[str, Any], named_styles: dict[str, Any]) -> int:
    text = str(block.get("text", ""))
    level = int(block.get("level", 1))
    if not 1 <= level <= 9:
        raise _DocxError(
            "INVALID_BLOCK",
            f"heading.level doit être entre 1 et 9, reçu : {level}",
            details={"field": "level"},
        )
    h = doc.add_heading(text, level=level)
    style_id = block.get("style")
    if style_id:
        if style_id not in named_styles:
            raise _DocxError(
                "STYLE_NOT_FOUND",
                f"Style '{style_id}' référencé mais absent de named_styles",
                details={"style_id": style_id},
            )
        _apply_paragraph_style(h, named_styles[style_id])
        return 1
    return 0


def _add_table_block(doc: Any, block: dict[str, Any], named_styles: dict[str, Any]) -> int:
    from docx.shared import Cm

    rows_spec = block.get("rows")
    if not isinstance(rows_spec, list) or not rows_spec:
        raise _DocxError("INVALID_BLOCK", "table.rows requis (liste non vide)", details={"field": "rows"})

    # Determine number of columns from first row
    first_row = rows_spec[0]
    if not isinstance(first_row, list) or not first_row:
        raise _DocxError("INVALID_BLOCK", "table.rows[0] doit être une liste non vide", details={"field": "rows[0]"})
    ncols = len(first_row)

    table = doc.add_table(rows=len(rows_spec), cols=ncols)
    styles_applied = 0

    # Fill cells
    for r_idx, row_data in enumerate(rows_spec):
        if not isinstance(row_data, list):
            raise _DocxError("INVALID_BLOCK", f"table.rows[{r_idx}] doit être une liste", details={"field": f"rows[{r_idx}]"})
        for c_idx, cell_data in enumerate(row_data):
            if c_idx >= ncols:
                break
            cell = table.cell(r_idx, c_idx)
            if isinstance(cell_data, str):
                cell.text = cell_data
            elif isinstance(cell_data, dict):
                cell.text = str(cell_data.get("text", ""))
                cell_style_id = cell_data.get("style")
                if cell_style_id:
                    if cell_style_id not in named_styles:
                        raise _DocxError(
                            "STYLE_NOT_FOUND",
                            f"Style '{cell_style_id}' référencé en cell mais absent de named_styles",
                            details={"style_id": cell_style_id},
                        )
                    for p in cell.paragraphs:
                        _apply_paragraph_style(p, named_styles[cell_style_id])
                    styles_applied += 1
                cell_align = cell_data.get("alignment")
                if cell_align:
                    for p in cell.paragraphs:
                        p.alignment = _str_to_alignment(cell_align)
                # Horizontal merge (gridSpan)
                merge = cell_data.get("merge")
                if isinstance(merge, dict):
                    cols_to_merge = int(merge.get("cols", 1))
                    if cols_to_merge > 1 and c_idx + cols_to_merge - 1 < ncols:
                        target_cell = table.cell(r_idx, c_idx + cols_to_merge - 1)
                        cell.merge(target_cell)
            else:
                cell.text = str(cell_data)

    # Apply column widths
    col_widths = block.get("column_widths_cm")
    if col_widths:
        if not isinstance(col_widths, list):
            raise _DocxError("INVALID_BLOCK", "column_widths_cm doit être une liste", details={"field": "column_widths_cm"})
        for c_idx, width in enumerate(col_widths):
            if c_idx >= ncols:
                break
            for row in table.rows:
                row.cells[c_idx].width = Cm(float(width))

    # Apply table style
    style_id = block.get("style")
    if style_id:
        if style_id not in named_styles:
            raise _DocxError(
                "STYLE_NOT_FOUND",
                f"Style '{style_id}' référencé sur table mais absent de named_styles",
                details={"style_id": style_id},
            )
        _apply_table_style(table, named_styles[style_id])
        styles_applied += 1

    return styles_applied


def _add_list_block(doc: Any, block: dict[str, Any], named_styles: dict[str, Any]) -> int:
    items = block.get("items")
    list_type = block.get("list_type", "bullet")
    if not isinstance(items, list) or not items:
        raise _DocxError("INVALID_BLOCK", "list.items requis (liste non vide)", details={"field": "items"})
    if list_type not in ("bullet", "number"):
        raise _DocxError(
            "INVALID_BLOCK",
            f"list_type invalide : {list_type!r}. Valeurs : ['bullet', 'number']",
            details={"field": "list_type"},
        )

    style_name = "List Bullet" if list_type == "bullet" else "List Number"
    style_id = block.get("style")
    styles_applied = 0

    for item in items:
        if isinstance(item, str):
            p = doc.add_paragraph(item, style=style_name)
        elif isinstance(item, dict):
            p = doc.add_paragraph(str(item.get("text", "")), style=style_name)
            # Nested items: flatten with indent (v0.1.0 limitation)
            sub_items = item.get("items")
            if sub_items:
                for sub in sub_items:
                    sub_text = sub if isinstance(sub, str) else str(sub.get("text", ""))
                    sub_p = doc.add_paragraph(f"    {sub_text}", style=style_name)
                    if style_id and style_id in named_styles:
                        _apply_paragraph_style(sub_p, named_styles[style_id])
                        styles_applied += 1
        else:
            p = doc.add_paragraph(str(item), style=style_name)

        if style_id:
            if style_id not in named_styles:
                raise _DocxError(
                    "STYLE_NOT_FOUND",
                    f"Style '{style_id}' référencé sur list mais absent de named_styles",
                    details={"style_id": style_id},
                )
            _apply_paragraph_style(p, named_styles[style_id])
            styles_applied += 1

    return styles_applied


def _add_image_block(doc: Any, block: dict[str, Any]) -> int:
    from docx.shared import Cm

    image_path = block.get("path")
    if not isinstance(image_path, str) or not image_path:
        raise _DocxError("INVALID_BLOCK", "image.path requis", details={"field": "path"})

    p = Path(image_path).expanduser().resolve()
    if not p.exists():
        raise _DocxError("IMAGE_NOT_FOUND", f"Image introuvable : {image_path}", details={"path": image_path})
    if p.suffix.lower() not in SUPPORTED_IMAGE_EXTS:
        raise _DocxError(
            "IMAGE_UNSUPPORTED",
            f"Format d'image non supporté : {p.suffix}. Valeurs : {list(SUPPORTED_IMAGE_EXTS)}",
            details={"ext": p.suffix},
        )

    width_cm = block.get("width_cm")
    height_cm = block.get("height_cm")

    kwargs: dict[str, Any] = {}
    if width_cm is not None:
        kwargs["width"] = Cm(float(width_cm))
    if height_cm is not None:
        kwargs["height"] = Cm(float(height_cm))

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(str(p), **kwargs)

    alignment = block.get("alignment")
    if alignment:
        paragraph.alignment = _str_to_alignment(alignment)

    return 0


def _add_page_break(doc: Any) -> None:
    from docx.enum.text import WD_BREAK

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ─── Style application ─────────────────────────────────────────────────────


def _apply_paragraph_style(paragraph: Any, style_spec: dict[str, Any]) -> None:
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_LINE_SPACING

    font_spec = style_spec.get("font")
    if font_spec:
        for run in paragraph.runs:
            _apply_font_to_run(run, font_spec)
        # If paragraph has no runs (empty), add a placeholder run with the font
        if not paragraph.runs:
            r = paragraph.add_run("")
            _apply_font_to_run(r, font_spec)

    para_spec = style_spec.get("paragraph")
    if para_spec:
        if not isinstance(para_spec, dict):
            raise _DocxError("INVALID_STYLE", "style.paragraph doit être un objet")
        pf = paragraph.paragraph_format
        if "alignment" in para_spec:
            paragraph.alignment = _str_to_alignment(para_spec["alignment"])
        if "space_before_pt" in para_spec:
            pf.space_before = Pt(float(para_spec["space_before_pt"]))
        if "space_after_pt" in para_spec:
            pf.space_after = Pt(float(para_spec["space_after_pt"]))
        if "line_spacing" in para_spec:
            pf.line_spacing = float(para_spec["line_spacing"])
        if "first_line_indent_cm" in para_spec:
            pf.first_line_indent = Cm(float(para_spec["first_line_indent_cm"]))
        if "left_indent_cm" in para_spec:
            pf.left_indent = Cm(float(para_spec["left_indent_cm"]))
        if "right_indent_cm" in para_spec:
            pf.right_indent = Cm(float(para_spec["right_indent_cm"]))


def _apply_font_to_run(run: Any, font_spec: dict[str, Any]) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX

    if not isinstance(font_spec, dict):
        raise _DocxError("INVALID_STYLE", "font doit être un objet")
    font = run.font
    if "name" in font_spec:
        font.name = font_spec["name"]
    if "size" in font_spec or "size_pt" in font_spec:
        size = font_spec.get("size_pt", font_spec.get("size"))
        font.size = Pt(float(size))
    if "bold" in font_spec:
        font.bold = bool(font_spec["bold"])
    if "italic" in font_spec:
        font.italic = bool(font_spec["italic"])
    if "underline" in font_spec:
        font.underline = bool(font_spec["underline"])
    if "strike" in font_spec:
        font.strike = bool(font_spec["strike"])
    if "color" in font_spec:
        hex_color = _validate_hex_color(font_spec["color"])
        font.color.rgb = RGBColor.from_string(hex_color[-6:])
    if "highlight" in font_spec:
        h = font_spec["highlight"]
        if isinstance(h, str):
            wd_color = _highlight_name_to_wd(h)
            if wd_color is not None:
                font.highlight_color = wd_color


_HIGHLIGHT_MAP = {
    "yellow": "YELLOW", "green": "GREEN", "cyan": "TURQUOISE",
    "magenta": "PINK", "blue": "BLUE", "red": "RED",
    "darkblue": "DARK_BLUE", "darkcyan": "TEAL", "darkgreen": "DARK_GREEN",
    "darkmagenta": "VIOLET", "darkred": "DARK_RED", "darkyellow": "DARK_YELLOW",
    "darkgray": "GRAY_50", "lightgray": "GRAY_25", "black": "BLACK", "white": "WHITE",
}


def _highlight_name_to_wd(name: str) -> Any | None:
    from docx.enum.text import WD_COLOR_INDEX

    key = name.lower().strip()
    if key not in _HIGHLIGHT_MAP:
        return None
    return getattr(WD_COLOR_INDEX, _HIGHLIGHT_MAP[key], None)


def _apply_table_style(table: Any, style_spec: dict[str, Any]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsmap

    table_spec = style_spec.get("table") or {}
    if not isinstance(table_spec, dict):
        raise _DocxError("INVALID_STYLE", "style.table doit être un objet")

    autofit = table_spec.get("autofit")
    if autofit is not None:
        table.autofit = bool(autofit)

    alignment = table_spec.get("alignment")
    if alignment:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        align_map = {"left": "LEFT", "center": "CENTER", "right": "RIGHT"}
        if alignment in align_map:
            table.alignment = getattr(WD_TABLE_ALIGNMENT, align_map[alignment])

    # Borders — apply to all cells (cleaner than table-level borders)
    borders = table_spec.get("borders")
    if borders and isinstance(borders, dict):
        _apply_table_borders(table, borders)

    # Cell shading
    shading = table_spec.get("cell_shading")
    if shading and isinstance(shading, dict):
        color = shading.get("color")
        if color:
            hex_color = _validate_hex_color(color)
            for row in table.rows:
                for cell in row.cells:
                    _set_cell_shading(cell, hex_color[-6:])


def _apply_table_borders(table: Any, borders_spec: dict[str, Any]) -> None:
    from docx.oxml import OxmlElement

    # Build tblBorders element
    tbl_pr = table._element.find(_qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)

    # Remove existing tblBorders
    existing = tbl_pr.find(_qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_borders = OxmlElement("w:tblBorders")
    valid_sides = ("top", "bottom", "left", "right", "inside_h", "inside_v")
    side_map = {"top": "top", "bottom": "bottom", "left": "left", "right": "right", "inside_h": "insideH", "inside_v": "insideV"}

    for side_name in valid_sides:
        spec = borders_spec.get(side_name)
        if not spec or not isinstance(spec, dict):
            continue
        border_elem = OxmlElement(f"w:{side_map[side_name]}")
        border_elem.set(_qn("w:val"), str(spec.get("style", "single")))
        size = spec.get("size_pt", 0.5)
        # OOXML border size is in eighths of a point
        border_elem.set(_qn("w:sz"), str(int(float(size) * 8)))
        color = spec.get("color", "000000")
        hex_color = _validate_hex_color(color)
        border_elem.set(_qn("w:color"), hex_color[-6:])
        tbl_borders.append(border_elem)

    tbl_pr.append(tbl_borders)


def _set_cell_shading(cell: Any, hex_color_rgb: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shd
    existing = tc_pr.find(_qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:val"), "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"), hex_color_rgb)
    tc_pr.append(shd)


# ─── Operation: append-section ─────────────────────────────────────────────


def _do_append_section(payload: dict[str, Any]) -> dict[str, Any]:
    path = _require_str(payload, "path")
    _validate_docx_path(path, must_exist=True)

    blocks = payload.get("blocks")
    if not isinstance(blocks, list) or not blocks:
        raise _DocxError("MISSING_FIELD", "Champ 'blocks' requis (liste non vide)", details={"field": "blocks"})

    named_styles = payload.get("named_styles") or {}
    section_break_before = bool(payload.get("section_break_before", False))
    new_section_setup = payload.get("new_section_setup") or {}

    from docx import Document
    from docx.enum.text import WD_BREAK

    try:
        doc = Document(path)
    except Exception as exc:
        raise _DocxError("PARSE_ERROR", f"Impossible de lire le .docx : {exc}") from exc

    if section_break_before:
        # Add page break + new section
        new_p = doc.add_paragraph()
        new_p.add_run().add_break(WD_BREAK.PAGE)
        new_section = doc.add_section()
        if new_section_setup:
            _apply_section_setup(new_section, new_section_setup, {})

    blocks_appended = 0
    for idx, block in enumerate(blocks):
        if not isinstance(block, dict):
            raise _DocxError(
                "INVALID_BLOCK",
                f"blocks[{idx}] doit être un objet",
                details={"field": f"blocks[{idx}]"},
            )
        _add_block(doc, block, named_styles)
        blocks_appended += 1

    try:
        doc.save(path)
    except Exception as exc:
        raise _DocxError("EXECUTION_FAILED", f"Sauvegarde impossible : {exc}") from exc

    return {
        "path": path,
        "blocks_appended": blocks_appended,
        "new_section_started": section_break_before,
    }


# ─── Operation: extract-tables ─────────────────────────────────────────────


def _do_extract_tables(payload: dict[str, Any]) -> dict[str, Any]:
    path = _require_str(payload, "path")
    _validate_docx_path(path, must_exist=True)
    include_headers = bool(payload.get("include_headers", True))

    from docx import Document

    try:
        doc = Document(path)
    except Exception as exc:
        raise _DocxError("PARSE_ERROR", f"Impossible de lire le .docx : {exc}") from exc

    tables_out: list[dict[str, Any]] = []
    for idx, tbl in enumerate(doc.tables):
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        entry: dict[str, Any] = {
            "index": idx,
            "rows": rows,
            "rows_count": len(rows),
            "cols_count": len(rows[0]) if rows else 0,
        }
        if include_headers and rows:
            entry["headers"] = rows[0]
        tables_out.append(entry)

    return {
        "path": path,
        "tables": tables_out,
        "total_tables": len(tables_out),
    }


# ─── Operation: render-from-template ───────────────────────────────────────


def _do_render_from_template(payload: dict[str, Any]) -> dict[str, Any]:
    template_path = _require_str(payload, "template_path")
    output_path = _require_str(payload, "output_path")
    _validate_docx_path(template_path, must_exist=True)
    _validate_docx_extension(output_path)

    context = payload.get("context")
    if not isinstance(context, dict):
        raise _DocxError(
            "MISSING_FIELD",
            "Champ 'context' requis (objet, peut être {})",
            details={"field": "context"},
        )

    strict_undefined = bool(payload.get("strict_undefined", True))
    overwrite = bool(payload.get("overwrite", False))

    out = Path(output_path).expanduser().resolve()
    if out.exists() and not overwrite:
        raise _DocxError(
            "OUTPUT_EXISTS",
            f"Fichier de sortie existe déjà (overwrite=false) : {out}",
            details={"output_path": str(out)},
        )
    out.parent.mkdir(parents=True, exist_ok=True)

    from docxtpl import DocxTemplate
    from jinja2 import StrictUndefined, Undefined
    from jinja2.sandbox import SandboxedEnvironment

    try:
        tpl = DocxTemplate(template_path)
    except Exception as exc:
        raise _DocxError("PARSE_ERROR", f"Impossible de lire le template : {exc}") from exc

    # docxtpl 0.18 : get_undeclared_template_variables() retourne TOUTES les
    # variables déclarées dans le template (pas seulement les non-fournies).
    # On fait la différence manuellement avec les clés du context.
    declared_variables: set[str] = set()
    try:
        declared_variables = tpl.get_undeclared_template_variables() or set()
    except Exception:
        declared_variables = set()
    undefined_variables = sorted(declared_variables - set(context.keys()))

    if strict_undefined and undefined_variables:
        raise _DocxError(
            "UNDEFINED_VARIABLES",
            f"Variables Jinja2 non fournies dans context : {undefined_variables}",
            details={"undefined": undefined_variables},
        )

    env = SandboxedEnvironment(undefined=StrictUndefined if strict_undefined else Undefined)

    try:
        tpl.render(context, jinja_env=env)
    except Exception as exc:
        raise _DocxError(
            "TEMPLATE_RENDER_ERROR",
            f"Erreur de rendu Jinja2 : {exc}",
            details={"reason": str(exc)},
        ) from exc

    try:
        tpl.save(str(out))
    except Exception as exc:
        raise _DocxError("EXECUTION_FAILED", f"Sauvegarde impossible : {exc}") from exc

    # Variables effectively used (those in context AND declared in template)
    variables_used = sorted(set(context.keys()) & declared_variables) if declared_variables else sorted(context.keys())

    return {
        "output_path": str(out),
        "template_path": template_path,
        "variables_used": variables_used,
        "undefined_variables": undefined_variables,
    }


# ─── Helpers ───────────────────────────────────────────────────────────────


def _require_str(payload: dict[str, Any], field: str) -> str:
    value = payload.get(field)
    if not isinstance(value, str) or not value:
        raise _DocxError(
            "MISSING_FIELD",
            f"Champ '{field}' requis (string non vide)",
            details={"field": field},
        )
    return value


def _validate_docx_extension(path: str) -> None:
    p = path.lower()
    if p.endswith(".doc") and not p.endswith(".docx"):
        raise _DocxError(
            "UNSUPPORTED_FORMAT",
            f"Format Word 97-2003 (.doc) non supporté : {path!r}",
            details={"path": path},
        )
    if p.endswith(".docm"):
        raise _DocxError(
            "UNSUPPORTED_FORMAT",
            f"Format avec macros (.docm) non supporté : {path!r}",
            details={"path": path},
        )
    if not p.endswith(".docx"):
        raise _DocxError(
            "UNSUPPORTED_FORMAT",
            f"Extension non reconnue : {path!r}. .docx requis.",
            details={"path": path},
        )


def _validate_docx_path(path: str, must_exist: bool = False) -> None:
    _validate_docx_extension(path)
    if must_exist:
        p = Path(path)
        if not p.exists():
            raise _DocxError("FILE_NOT_FOUND", f"fichier introuvable : {path}")
        size = p.stat().st_size
        if size > MAX_FILE_BYTES:
            raise _DocxError(
                "TOO_LARGE",
                f"Fichier trop volumineux ({size} > {MAX_FILE_BYTES} bytes)",
                details={"size_bytes": size, "limit_bytes": MAX_FILE_BYTES},
            )


def _validate_hex_color(color: Any) -> str:
    if not isinstance(color, str):
        raise _DocxError(
            "INVALID_STYLE",
            f"Couleur doit être hex string, reçu : {type(color).__name__}",
        )
    c = color.lstrip("#").upper()
    if len(c) == 6 and all(ch in "0123456789ABCDEF" for ch in c):
        return "FF" + c
    if len(c) == 8 and all(ch in "0123456789ABCDEF" for ch in c):
        return c
    raise _DocxError(
        "INVALID_STYLE",
        f"Couleur hex invalide : {color!r} (attendu RRGGBB ou AARRGGBB)",
    )


def _str_to_alignment(name: Any) -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not isinstance(name, str):
        raise _DocxError(
            "INVALID_TYPE",
            f"alignment doit être string, reçu : {type(name).__name__}",
        )
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    key = name.lower().strip()
    if key not in align_map:
        raise _DocxError(
            "INVALID_TYPE",
            f"alignment invalide : {name!r}. Valeurs : {list(align_map)}",
            details={"field": "alignment"},
        )
    return align_map[key]


def _alignment_to_str(align: Any) -> str:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rev = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return rev.get(align, "left")


def _emu_to_cm(emu: Any) -> float | None:
    if emu is None:
        return None
    return float(emu) / 360000.0


def _qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def _json_default(o: Any) -> Any:
    if isinstance(o, (datetime.datetime, datetime.date)):
        return o.isoformat()
    raise TypeError(f"Object of type {type(o).__name__} is not JSON serializable")


# (The ``agent`` module-level singleton is exposed automatically by the
# ``@agent`` decorator.)
